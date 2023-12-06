from selenium import webdriver
from selenium.webdriver.remote.webelement import WebElement
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from time import sleep
from docx import Document
from docx.shared import Inches
from docx.shared import Pt, RGBColor
import requests
import os
from PIL import Image

class TayrBot:
    def __init__(self) -> None:
        self.url = 'https://signin.valuegaia.com.br/?provider=imob'
        self.username = 'Wanderley.grupotrevo@gmail.com'
        self.password = 'Wa2022.2022'
        self.driver = webdriver.Chrome()
        self.driver.get(self.url)


    def find_login_inputs(self) -> tuple[WebElement, WebElement, WebElement]:
        email_input = WebDriverWait(self.driver, 10).until(EC.presence_of_element_located((By.ID, "email")))
        pwd_input = WebDriverWait(self.driver, 10).until(EC.presence_of_element_located((By.ID, "password")))
        login_btn = WebDriverWait(self.driver, 10).until(EC.presence_of_element_located((By.ID, "enter-login")))
    
        return email_input, pwd_input, login_btn
    
    def interact_with_login_inputs(self, email_input:str, pwd_input:str, login_btn:str) -> None:
        email_input.send_keys(self.username)
        pwd_input.send_keys(self.password)

        sleep(20) #Time to answer captcha
        login_btn.click()
        sleep(10)

    def find_menu_icon_and_move_through(self) -> None:
        icon_menu = WebDriverWait(self.driver, 10).until(EC.element_to_be_clickable((By.ID, "step-toggle-menu")))
        icon_menu.click()
        sleep(0.5)
        id_imoveis = WebDriverWait(self.driver, 10).until(EC.element_to_be_clickable((By.ID, "imoveis")))
        id_imoveis.click()
        sleep(0.5)
        id_imoveis = WebDriverWait(self.driver, 10).until(EC.element_to_be_clickable((By.ID, "imoveis_imoveis")))
        id_imoveis.click()
        sleep(5)

    def switch_to_iframe_and_submit(self) -> None:
        self.driver.switch_to.frame("iframePage")
        button = WebDriverWait(self.driver, 10).until(EC.presence_of_element_located((By.ID, 'fieldsContainer')))
        button.submit()
        sleep(3)

    def delete_files_in_folder(self, folder_path:str) -> None:
        if os.path.isdir(folder_path):
            files = os.listdir(folder_path)

            for file in files:
                full_path = os.path.join(folder_path, file)
                try:
                    if os.path.isfile(full_path):
                        os.remove(full_path)
                    else:
                        print(f"{full_path} is not a file.")
                except Exception as e:
                    print(f"Error deleting {full_path}: {e}")

    def clean_title(self, title:str) -> str:
        invalid_chars = set(r'\/:*?"<>|')
        cleaned_title = ''.join(c for c in title if c not in invalid_chars)
        
        return cleaned_title

    def add_image_to_doc(self, doc, image_path:str, width:float = 6, height:float = 6) -> None:
        try:
            # Open and save the image using Pillow to correct the header
            img = Image.open(image_path)
            img.save(image_path, format="PNG")

            # Add the image to the document
            doc.add_picture(image_path, width=width, height=height)
        except Exception as e:
            print(f"Error processing image {image_path}: {e}")

    def create_doc(self, building, name, price, description, title, address, rent, phone) -> None:
        doc = Document()
        
        title_doc = doc.add_heading(f"{title}", level=1)
        title_doc.alignment = 1
        title_doc.runs[0].font.size = Pt(24)
        title_doc.runs[0].font.color.rgb = RGBColor(0, 0, 0)

        doc.add_paragraph('')

        description_doc = doc.add_paragraph(f'{description}')
        description_doc.runs[0].font.size = Pt(18)
        doc.add_paragraph('')

        building_doc = doc.add_paragraph(f"Construção:\n {building}")
        building_doc.runs[0].font.size = Pt(18)
        doc.add_paragraph('')

        address_doc = doc.add_paragraph(f"Endereço: {address}")
        address_doc.runs[0].font.size = Pt(18)
        doc.add_paragraph('')

        price_doc = doc.add_paragraph(f"{price}, {rent}")
        price_doc.runs[0].font.size = Pt(18)
        doc.add_paragraph('')

        name_doc = doc.add_paragraph(f"Nome do proprietário: {name}")
        name_doc.runs[0].font.size = Pt(18)
        doc.add_paragraph('')

        phone_doc = doc.add_paragraph(f"Contato: {phone}")
        phone_doc.runs[0].font.size = Pt(18)
        doc.add_paragraph('')

        photos = doc.add_paragraph("Fotos")
        photos.runs[0].font.size = Pt(24)
        photos.alignment = 1

        image_folder = "D:\\Projects\\tayrbot\\docsimage"

        for image in os.listdir(image_folder):
            image_path = os.path.join(image_folder, image)
            self.add_image_to_doc(doc, image_path, width=Inches(6), height=Inches(6))

        doc.save(f'D:\\Projects\\tayrbot\\docs\\{title}.docx')

    def get_image(self) -> str:
        self.driver.switch_to.frame("propertyRecordIframe")
        image_folder = "D:\\Projects\\tayrbot\\docsimage"

        div_photo = WebDriverWait(self.driver, 10).until(EC.element_to_be_clickable((By.ID, "fotoMin")))
        div_photo.click()

        images = WebDriverWait(self.driver, 10).until(EC.presence_of_all_elements_located((By.TAG_NAME, "img")))
        
        try:
            len_images = WebDriverWait(self.driver, 10).until(EC.presence_of_element_located((By.CSS_SELECTOR, "#PhotoViewerTime")))
            len_images = len_images.text
        except:
            self.driver.switch_to.default_content()
            self.driver.switch_to.frame("iframePage")
            
            return image_folder
        
        slash_index = len_images.find("/")+1
        len_images = len_images[slash_index:]
        len_images = int(len_images)

        count = 0
        
        for image in images:
            count += 1
            image_link = image.get_attribute("src")

            if image_link.endswith("w100-h75"):
                image_link = image_link.replace("w100-h75","w1280-h720")
            
            response = requests.get(image_link)

            if count > 5:
                break
            
            elif response.status_code == 200:
                if not os.path.exists(image_folder):
                    os.makedirs(image_folder)

                output = f"D:\\Projects\\tayrbot\\docsimage\\{count}.jpeg"
                with open(output, 'wb') as file:
                    file.write(response.content)

                if len_images == 1:
                    break
            else:
                print(f"Failed to download the image {count}.") 

        self.driver.switch_to.default_content()
        self.driver.switch_to.frame("iframePage")
        
        return image_folder
            
    def get_contact(self) -> tuple[str, str]:
        self.driver.switch_to.frame("propertyRecordIframe")

        try:
            confidencial = WebDriverWait(self.driver, 10).until(EC.element_to_be_clickable((By.LINK_TEXT, "Confidencial")))
            confidencial.click()
        except:
            self.driver.switch_to.default_content()
            self.driver.switch_to.frame("iframePage")
            return "S/N", "S/N"

        try:
            name_element = WebDriverWait(self.driver, 10).until(EC.presence_of_element_located((By.CSS_SELECTOR, "#tabs-6 > div:nth-child(3) > span")))
            name = name_element.text
        except:
            name = "S/N"

        try:
            phone_element = WebDriverWait(self.driver, 10).until(EC.presence_of_element_located((By.CSS_SELECTOR, "#tabs-6 > div:nth-child(5) > span > div")))
            phone = phone_element.text
        except:
            phone = "S/N"

        self.driver.switch_to.default_content()
        self.driver.switch_to.frame("iframePage")

        return name, phone

    def get_price(self) -> tuple[str, str]:
        self.driver.switch_to.frame("propertyRecordIframe")

        price = ""
        rent = ""

        data_tab = WebDriverWait(self.driver, 10).until(EC.element_to_be_clickable((By.LINK_TEXT, "Dados prim")))
        sleep(0.1)
        data_tab.click()
        itens = WebDriverWait(self.driver, 10).until(EC.presence_of_all_elements_located((By.CLASS_NAME, "itens")))
        for item in itens:
            item = item.text
            
            if item.startswith("$ Venda") and item != "":
                price = item
            if item.startswith("$ Locação") and item != "":
                rent = item
        
        self.driver.switch_to.default_content()
        self.driver.switch_to.frame("iframePage")
        return price, rent

    def get_address(self) -> str:
        self.driver.switch_to.frame("propertyRecordIframe")
        address = WebDriverWait(self.driver, 10).until(EC.presence_of_element_located((By.TAG_NAME, "address")))
        address = address.text
        self.driver.switch_to.default_content()
        self.driver.switch_to.frame("iframePage")

        if address == "":
            return "S/N"
        
        return address

    def get_description(self) -> str:
        self.driver.switch_to.frame("propertyRecordIframe")
        ad_tab = WebDriverWait(self.driver, 10).until(EC.element_to_be_clickable((By.LINK_TEXT, "Anúncio")))
        ad_tab.click()
        description = WebDriverWait(self.driver, 10).until(EC.presence_of_element_located((By.CLASS_NAME, "txTitulo")))
        description = description.text
        
        self.driver.switch_to.default_content()
        self.driver.switch_to.frame("iframePage")
        
        return description

    def get_title(self) -> str:
        self.driver.switch_to.frame("propertyRecordIframe")
        title = WebDriverWait(self.driver, 10).until(EC.element_to_be_clickable((By.TAG_NAME, "h1")))
        title = title.text
        
        self.driver.switch_to.default_content()
        self.driver.switch_to.frame("iframePage")

        return title

    def get_building(self) -> str:
        self.driver.switch_to.frame("propertyRecordIframe")
        building = WebDriverWait(self.driver, 10).until(EC.presence_of_element_located((By.CSS_SELECTOR, "div[class='column4'] pre")))
        building = building.text

        self.driver.switch_to.default_content()
        self.driver.switch_to.frame("iframePage")

        return building

    def iterate_cards(self) -> list[WebElement]:
        cards = WebDriverWait(self.driver, 10).until(EC.presence_of_all_elements_located((By.CLASS_NAME, "imgImoveis")))
        return cards

    def move_through_cards(self, count, is_finished, cards) -> int:
        while is_finished == False:
            try:
                for card in cards:
                    card.click()
                    sleep(0.5)

                    building = self.get_building()
                    name, phone = self.get_contact()
                    price, rent = self.get_price()
                    description = self.get_description()
                    title = self.get_title()
                    address = self.get_address()
                    title = self.clean_title(title=title)
                    image_folder = self.get_image()

                    close_button = WebDriverWait(self.driver, 10).until(EC.presence_of_element_located((By.CLASS_NAME, "closePropertyRecord")))
                    close_button.click()
                    sleep(0.5)
                    count += 1

                    self.create_doc(building=building, name=name, phone=phone, address=address, description=description, price=price, rent=rent, title=title)
                    self.delete_files_in_folder(image_folder)

                    print(f"Arquivo {count} concluído.")

                    if card == cards[-1]:
                        try:
                            button = WebDriverWait(self.driver, 10).until(EC.element_to_be_clickable((By.CLASS_NAME, "btn-md")))
                            button.click()
                        except Exception:
                            print(f"Cards finished, exiting software...")
                        is_finished = True
            
            except Exception as e:
                print(f"Error: {e}")
                continue

        return count
    
    def close_ad(self) -> None:
        popup = WebDriverWait(self.driver, 10).until(EC.presence_of_element_located((By.ID, "beamerAnnouncementPopup")))

        if popup:
            self.driver.switch_to.frame("beamerAnnouncementPopup")
            close_button = WebDriverWait(self.driver, 10).until(EC.element_to_be_clickable((By.CLASS_NAME, "popupClose")))
            close_button.click()
            self.driver.switch_to.default_content()

    def main(self) -> None:
        cards = self.iterate_cards()

        count = 0
        is_finished = False
        
        while True:
            count = self.move_through_cards(count, is_finished, cards)
            sleep(1)

            try:
                cards = self.iterate_cards()[count:]
            except IndexError:
                break

            if not cards:
                break
            
tayrbot = TayrBot()

email_input, pwd_input, login_btn = tayrbot.find_login_inputs()
tayrbot.interact_with_login_inputs(email_input, pwd_input, login_btn)

tayrbot.close_ad()
tayrbot.find_menu_icon_and_move_through()
tayrbot.switch_to_iframe_and_submit()

tayrbot.main()
tayrbot.driver.quit()